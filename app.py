# -*- coding: utf-8 -*-
from os import path

from difflib import SequenceMatcher
from tkinter import Tk, Menu, LabelFrame, Frame, Label, Button, Text, Scrollbar, RIGHT, LEFT, BOTTOM, TOP, BOTH, X, Y, Toplevel, END
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter.messagebox import showwarning, showinfo
from tkinter.ttk import Combobox
import tkinter.font as tkFont
import docx

def fontScale(scale=12, fontfamily='Times'):
    font = tkFont.Font(size=scale, family=fontfamily)
    return font
def normalize(string):
	normal = string.lower()
	normal = string.replace('\n', '')
	return normal

class App:
	def __init__(self, root):
		self.source = []
		filenames = []
		symb = []
		report_summary = []
		symbol_cap = 25000
		def clearEdit():
			edit.delete(1.0, END)
		def clearReport():
			self.source.clear()
			filenames.clear()
			report['state']='normal'
			report.delete(1.0, END)
			report['state']='disabled'
			status_label['text']=f'Файлов в базе сравнения: {len(self.source)}'
		def clearAll():
			clearEdit()
			clearReport()

		def statusUpdate(current, symbols):
			status_label['text']=f'Файлов в базе сравнения: {len(self.source)} | Текущий файл: {current} | Количество символов: {symbols}'

		def addSource():
				file = askopenfilename(filetypes=(("TXT файлы", "*.txt"),("DOCX файлы", "*.docx")))
				size = len(file)
				if file[size-3:] == 'txt':
					with open(file, 'r', encoding="utf8") as source_file:
						source = source_file.read()
						if len(source) > symbol_cap: return showWigdetHighSymbolsAmount()
						else:
							report['state']='normal'
							report.delete(1.0, END)
							report.insert('end', source)
							report['state']='disabled'
							self.source.append(source)
							filename = path.basename(file)
							filenames.append(filename)
							symb.append(len(source))
							statusUpdate(filename, len(source))
				elif file[size-4:] == 'docx':
					doc = docx.Document(file)
					source_list = []
					for i in range(len(doc.paragraphs)):
						source_list.append(doc.paragraphs[i].text)
					source = ''.join(source_list)
					if len(source) > symbol_cap: return showWigdetHighSymbolsAmount()
					else:
						report['state']='normal'
						report.delete(1.0, END)
						report.insert('end', source)
						report['state']='disabled'
						self.source.append(source)
						filename = path.basename(file)
						filenames.append(filename)
						symb.append(len(source))
						statusUpdate(filename, len(source))
				elif file == '': pass
				else: return showWidgetUnsupportedFormat()

		def plagiarismCheckByDifflib():
			searches = 0
			string1 = edit.get(1.0, END)
			checking_string = normalize(string1)
			if len(string1) == 1:
				return showWidgetEmptyEdit()
			elif len(self.source) == 0:
				return showWidgetEmptyReport()
			else:
				for i in range(len(self.source)):
					source_string = normalize(self.source[i])
					searches += 1
					if checking_string == source_string:
						report['state']='normal'
						report.delete(1.0, END)
						report.insert('end', source_string)
						report['state']='disabled'
						statusUpdate(filenames[i], symb[i])
						report_summary.append('0')
						report_summary.append(filenames[i])
						showinfo(title=f'Проверка №{searches}', message=f'Уникальность: 0%\nИсточник: {filenames[i]}')
					else:
						check = SequenceMatcher(None, checking_string, source_string).ratio() * 100
						unique = 100-check
						report['state']='normal'
						report.delete(1.0, END)
						report.insert('end', source_string)
						report['state']='disabled'
						statusUpdate(filenames[i], symb[i])
						report_summary.append(round(unique, 1))
						report_summary.append(filenames[i])
						showinfo(title=f'Проверка №{searches}', message=f'Уникальность: {round(unique, 1)}%\nИсточник: {filenames[i]}')

		def onOpen():
			file = askopenfilename(filetypes=(("TXT файлы", "*.txt"),("DOCX файлы", "*.docx")))
			size = len(file)
			if file[size-3:] == 'txt':
				with open(file, 'r', encoding="utf8") as checking:
					text = checking.read()
					if len(text) > symbol_cap: return showWigdetHighSymbolsAmount()
					else:
						edit.delete(1.0, END)
						edit.insert(1.0, text)
			elif file[size-4:] == 'docx':
				doc = docx.Document(file)
				checking_list = []
				for i in range(len(doc.paragraphs)):
					checking_list.append(doc.paragraphs[i].text)
				text = ''.join(checking_list)
				if len(text) > symbol_cap: return showWigdetHighSymbolsAmount()
				else:
					edit.delete(1.0, END)
					edit.insert(1.0, text)
			elif file == '': pass
			else: return showWidgetUnsupportedFormat()

		def saveReport():
			dlg = asksaveasfilename(filetypes=(("TXT файлы", "*.txt"),("All files", "*.*")))
			uniques = []
			files = []
			string = ''
			if not report_summary:
				return showWidgetEmptyReportSummary()
			elif dlg == '': pass
			else:
				string = ''
				for i in range(0, len(report_summary), 2):
					uniques.append(f'Уникальность: {str(report_summary[i])} ')
				for i in range(1, len(report_summary), 2):
					files.append(f'Файл-источник: {str(report_summary[i])}\n')
				for i in range(0, len(uniques)):
					string += uniques[i]
					string += files[i]
				with open(f'{dlg}.txt', "w", encoding='utf8') as file:
					file.write(string)
				showWidgetFileSaved()

		def howToUse():
			info = showinfo(title='Краткое руководство по использованию программы', message='Перед проверкой на плагиат необходимо импортировать источник (с которым будет проводиться сравнение) при помощи кнопки "Добавить источник".\nДобавить текст в редактор самостоятельно, или при помощи команды меню "Файл-открыть".\nЗапустить проверку при помощи кнопки "Начать проверку".\nРезультат можно сохранить в текстовый файл при помощи команды меню "Файл-Сохранить отчет".')

		def aboutPlagiarism():
			info = showinfo(title='Scripture v1.0', message='Программа для проверки на плагиат среди нескольких файлов.\nСоздана для дипломного проекта Донецкого техникума промышленной автоматики имени А.В. Захарченко')

		def showWidgetUnsupportedFormat():
			warning = showwarning(title='Неверный формат', message='Выбранный формат файла не поддерживается, выберете любой из следующих: .txt, .docx')
		def showWidgetEmptyReport():
			warning = showwarning(title='Отсутствует тест в базе сравнения', message='Добавьте текст в базу сравнения при помощи кнопки "Добавить источник"')
		def showWidgetEmptyEdit():
			warning = showwarning(title='Отсутствует тест для сравнения', message='Добавьте текст в редактор самостоятельно или при помощи меню "Файл" - "Открыть"')
		def showWidgetEmptyReportSummary():
			warning = showwarning(title='Проверка не проведена', message='Перед созданием отчёта необходимо провести проверку')
		def showWidgetFileSaved():
			info = showinfo(title='Сохранение завершено', message='Файл успешно сохранён')
		def showWigdetHighSymbolsAmount():
			warning = showwarning(title='Файл слишком большой', message=f'Импортирование файлов с более чем {symbol_cap} символов может вызвать зависания, выберете файл меньшего размера')
		root.title("Scripture - Антиплагиат")
		root.iconbitmap("icon.ico")
		width=1200
		height=600
		screenwidth = root.winfo_screenwidth()
		screenheight = root.winfo_screenheight()
		alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
		root.geometry(alignstr)
		root.resizable(width=False, height=False)

		menubar = Menu(root, bg='#d7f3f7')
		root['menu'] = menubar
		root['bg'] = '#d7f3f7'

		menu_file = Menu(menubar)
		menu_about = Menu(menubar)
		menubar.add_cascade(menu=menu_file, label='Файл')
		menubar.add_cascade(menu=menu_about, label='Справка')
		menu_file.add_command(label="Открыть", command=onOpen)
		menu_file.add_command(label="Сохранить отчет", command=saveReport)
		menu_about.add_command(label="Руководство по использованию", command=howToUse)
		menu_about.add_command(label="О программе", command=aboutPlagiarism)

		font = fontScale(12)
		status_bar = Frame(root, height=25, width=width, bg='#2bc2d9')
		status_bar.pack(side=BOTTOM)

		buttons_frame = Frame(root, width=160)
		buttons_frame.pack(side=LEFT)
		add_file = Button(buttons_frame, font=fontScale(14, 'Segoe'), width=16, height=1, text="Добавить источник", command=addSource)
		simple_check_button = Button(buttons_frame, font=fontScale(14, 'Segoe'), width=16, height=1, text="Начать проверку", command=plagiarismCheckByDifflib)
		clear_edit = Button(buttons_frame, font=fontScale(14, 'Segoe'), width=16, height=1, text="Очистить редактор", command=clearEdit)
		clear_report = Button(buttons_frame, font=fontScale(14, 'Segoe'), width=16, height=1, text="Очистить источники", command=clearReport)
		clear_all = Button(buttons_frame, font=fontScale(14, 'Segoe'), width=16, height=1, text="Очистить всё", command=clearAll)
		add_file.pack(side=TOP, padx=2, pady=3)
		simple_check_button.pack(side=TOP, padx=2, pady=3)
		clear_edit.pack(side=TOP, padx=2, pady=3)
		clear_report.pack(side=TOP, padx=2, pady=3)
		clear_all.pack(side=TOP, padx=2, pady=3)

		edit_frame = LabelFrame(root, font=fontScale(13, 'Segoe'), text="Редактор текста", width=450, height=550, bg='#d7f3f7')
		edit_frame.pack(side=LEFT, padx=5)
		edit_scroll = Scrollbar(edit_frame)
		edit_scroll.pack(side=RIGHT, fill=Y)
		edit = Text(edit_frame, font=font, width=45, height=550, yscrollcommand=edit_scroll.set)
		edit.pack()
		edit_scroll.config( command = edit.yview )

		report_frame = LabelFrame(root, font=fontScale(13, 'Segoe'), text="Источник для проверки", width=450, height=550, bg='#d7f3f7')
		report_frame.pack(side=LEFT, padx=5)
		report_scroll = Scrollbar(report_frame)
		report_scroll.pack(side=RIGHT, fill=Y)
		report = Text(report_frame, font=font, width=450, height=550, state='disabled', yscrollcommand=report_scroll.set)
		report.pack()
		report_scroll.config( command = report.yview )

		status_label = Label(status_bar, font=font, text="Файлов в базе сравнения: 0", bg='#d7f3f7')
		status_label.pack(side=TOP)
		#report_label = Label(root, text="Журнал проверки", bg='#ade7f0').place(relx=.455, rely=.01)
		#side=RIGHT, padx=10, pady=30
		
		#edit_label = Label(root, text="Редактор текста", bg='#ade7f0').place(relx=.155, rely=.01)
		#side=RIGHT, padx=10, pady=30
		
if __name__ == "__main__":
	root = Tk()
	app = App(root)
	root.mainloop()
